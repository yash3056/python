from PIL import Image
import fitz  # PyMuPDF
import pytesseract
import tabula
from docx import Document

def pdf_to_word_with_ocr_and_tables(pdf_file, word_file):
    # Open the PDF file
    pdf_document = fitz.open(pdf_file)

    # Create a new Word document
    doc = Document()

    # Extract text using OCR and append to a list
    ocr_text = []
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        ocr_text.append(text)

    # Extract tables from the PDF
    tables = tabula.read_pdf(pdf_file, pages='all', multiple_tables=True)

    # Loop through OCR text and insert into the Word document
    for text in ocr_text:
        doc.add_paragraph(text)

    # Loop through extracted tables and insert into the Word document
    for table in tables:
        if table is None:
            continue
        doc.add_table(rows=len(table), cols=len(table.columns))
        for i, row in enumerate(table.values):
            for j, cell in enumerate(row):
                doc.tables[-1].cell(i, j).text = str(cell)
        doc.add_paragraph()  # Add a paragraph after each table for spacing

    # Save the Word document
    doc.save(word_file)
    print(f"PDF converted with OCR and tables, and saved as '{word_file}'")

# Example usage
pdf_file = 'sample.pdf'
word_file = 'output.docx'
pdf_to_word_with_ocr_and_tables(pdf_file, word_file)